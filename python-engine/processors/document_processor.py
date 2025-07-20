import docx
from bs4 import BeautifulSoup
import markdown
import chardet
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass
import logging

@dataclass
class DocumentExtractionResult:
    text: str
    metadata: Dict
    structure: Dict
    extraction_method: str
    quality_score: float
    encoding: Optional[str] = None

class DocumentProcessor:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
    def process_document(self, file_path: str) -> DocumentExtractionResult:
        """Process any document type"""
        file_path = Path(file_path)
        mime_type, _ = mimetypes.guess_type(str(file_path))
        
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._process_docx(file_path)
        elif mime_type in ['text/html', 'application/xhtml+xml']:
            return self._process_html(file_path)
        elif file_path.suffix.lower() in ['.md', '.markdown']:
            return self._process_markdown(file_path)
        elif mime_type and mime_type.startswith('text/'):
            return self._process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
            
    def _process_docx(self, file_path: Path) -> DocumentExtractionResult:
        """Process DOCX files"""
        doc = docx.Document(str(file_path))
        
        # Extract text
        text_parts = []
        for paragraph in doc.paragraphs:
            text_parts.append(paragraph.text)
            
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                text_parts.append(" | ".join(row_text))
                
        full_text = "\n".join(text_parts)
        
        # Extract metadata
        props = doc.core_properties
        metadata = {
            'title': props.title or '',
            'author': props.author or '',
            'subject': props.subject or '',
            'keywords': props.keywords or '',
            'created': str(props.created) if props.created else '',
            'modified': str(props.modified) if props.modified else '',
            'last_modified_by': props.last_modified_by or '',
            'revision': props.revision or '',
            'category': props.category or '',
            'comments': props.comments or '',
        }
        
        # Analyze structure
        structure = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'sections': len(doc.sections),
            'has_headers': any(p.style.name.startswith('Heading') for p in doc.paragraphs),
            'has_tables': len(doc.tables) > 0,
        }
        
        quality_score = min(1.0, len(full_text.strip()) / 100)
        
        return DocumentExtractionResult(
            text=full_text,
            metadata=metadata,
            structure=structure,
            extraction_method='docx',
            quality_score=quality_score
        )
        
    def _process_html(self, file_path: Path) -> DocumentExtractionResult:
        """Process HTML files"""
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            
        # Detect encoding
        encoding_result = chardet.detect(raw_data)
        encoding = encoding_result['encoding'] or 'utf-8'
        
        html_content = raw_data.decode(encoding, errors='ignore')
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
            
        # Extract text
        text = soup.get_text()
        
        # Clean up text
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        # Extract metadata
        metadata = {}
        if soup.title:
            metadata['title'] = soup.title.string or ''
            
        # Meta tags
        for meta in soup.find_all('meta'):
            name = meta.get('name') or meta.get('property')
            content = meta.get('content')
            if name and content:
                metadata[name] = content
                
        # Analyze structure
        structure = {
            'headings': {f'h{i}': len(soup.find_all(f'h{i}')) for i in range(1, 7)},
            'paragraphs': len(soup.find_all('p')),
            'links': len(soup.find_all('a')),
            'images': len(soup.find_all('img')),
            'tables': len(soup.find_all('table')),
            'lists': len(soup.find_all(['ul', 'ol'])),
        }
        
        quality_score = min(1.0, len(text.strip()) / 200)
        
        return DocumentExtractionResult(
            text=text,
            metadata=metadata,
            structure=structure,
            extraction_method='html',
            quality_score=quality_score,
            encoding=encoding
        )
        
    def _process_markdown(self, file_path: Path) -> DocumentExtractionResult:
        """Process Markdown files"""
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            
        # Detect encoding
        encoding_result = chardet.detect(raw_data)
        encoding = encoding_result['encoding'] or 'utf-8'
        
        md_content = raw_data.decode(encoding, errors='ignore')
        
        # Parse markdown
        md = markdown.Markdown(extensions=['meta', 'toc', 'tables', 'fenced_code'])
        html = md.convert(md_content)
        
        # Extract plain text
        soup = BeautifulSoup(html, 'html.parser')
        text = soup.get_text()
        
        # Extract metadata from frontmatter
        metadata = getattr(md, 'Meta', {})
        metadata = {k: v[0] if isinstance(v, list) and len(v) == 1 else v 
                   for k, v in metadata.items()}
        
        # Analyze structure
        lines = md_content.split('\n')
        structure = {
            'headings': {f'h{i}': md_content.count('#' * i + ' ') for i in range(1, 7)},
            'code_blocks': md_content.count('```'),
            'inline_code': md_content.count('`') - md_content.count('```') * 6,
            'links': md_content.count('['),
            'images': md_content.count('!['),
            'tables': md_content.count('|'),
            'lists': sum(1 for line in lines if line.strip().startswith(('-', '*', '+'))),
            'numbered_lists': sum(1 for line in lines if line.strip() and line.strip()[0].isdigit()),
        }
        
        quality_score = min(1.0, len(text.strip()) / 200)
        
        return DocumentExtractionResult(
            text=text,
            metadata=metadata,
            structure=structure,
            extraction_method='markdown',
            quality_score=quality_score,
            encoding=encoding
        )
        
    def _process_text(self, file_path: Path) -> DocumentExtractionResult:
        """Process plain text files"""
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            
        # Detect encoding
        encoding_result = chardet.detect(raw_data)
        encoding = encoding_result['encoding'] or 'utf-8'
        confidence = encoding_result['confidence']
        
        text = raw_data.decode(encoding, errors='ignore')
        
        # Basic metadata
        metadata = {
            'encoding': encoding,
            'encoding_confidence': confidence,
            'file_size': len(raw_data),
        }
        
        # Analyze structure
        lines = text.split('\n')
        structure = {
            'lines': len(lines),
            'empty_lines': sum(1 for line in lines if not line.strip()),
            'avg_line_length': sum(len(line) for line in lines) / max(len(lines), 1),
            'max_line_length': max(len(line) for line in lines) if lines else 0,
            'paragraphs': len([p for p in text.split('\n\n') if p.strip()]),
        }
        
        quality_score = min(1.0, confidence * len(text.strip()) / 200)
        
        return DocumentExtractionResult(
            text=text,
            metadata=metadata,
            structure=structure,
            extraction_method='text',
            quality_score=quality_score,
            encoding=encoding
        )